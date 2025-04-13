#!/usr/bin/env python3
import sys, os, subprocess
from PyQt5 import QtWidgets, QtGui, QtCore
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
import pandas as pd
from itertools import combinations
from collections import defaultdict
from copy import deepcopy
import networkx as nx


# -------------------------- Глобальные переменные -------------------------- #
way = set()
LOGS = []
LAST_DF_UPPER = None
LAST_AUTOMATA = None
LAST_MAX_COVER = None
LAST_MIN_COVER = None
LAST_BIN_MATRIX = None


def log_msg(message: str):
    LOGS.append(message)
    print(message)


# -------------------------- Рисование покрытий с наложением -------------------------- #
def draw_coverings_with_overlap(coverings, identifier="coverings"):
    log_msg(f"Начало рисования покрытий ({identifier}) c учётом частичного перекрытия.")
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_title(f"Покрытия ({identifier.upper()})", fontsize=14)
    ax.axis('off')

    overlap_shift = 0.07
    separate_shift = 0.22
    x_current = 0.1
    prev_block = None

    for idx, block in enumerate(coverings):
        text = ", ".join(block)
        width = 0.15 + 0.03 * len(text)
        height = 0.3

        if idx > 0 and prev_block:
            if set(block).intersection(set(prev_block)):
                x_current += overlap_shift
            else:
                x_current += separate_shift

        ellipse = patches.Ellipse((x_current, 0.5), width, height,
                                  edgecolor='blue', facecolor='lightblue', lw=2)
        ax.add_patch(ellipse)
        ax.text(x_current, 0.5, text, fontsize=12,
                ha='center', va='center', color='darkblue', fontweight='bold')
        log_msg(f"Блок {idx + 1}: [{text}], x={x_current:.2f}, пересечение c пред. блоком: "
                f"{bool(set(block).intersection(set(prev_block))) if prev_block else 'N/A'}")

        prev_block = block

    plt.tight_layout()
    output_path = f"coverings_{identifier}.png"
    plt.savefig(output_path)
    log_msg(f"Покрытия c наложением сохранены в файл: {output_path}")
    plt.close(fig)


def draw_comparison_coverings(max_cover, min_cover):
    log_msg("Построение сравнительного отображения покрытий.")
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

    ax1.set_title("Максимальное покрытие", fontsize=12)
    ax1.axis('off')
    overlap_shift = 0.07
    separate_shift = 0.22
    x_current = 0.1
    prev_block = None
    for idx, block in enumerate(max_cover):
        text = ", ".join(block)
        width = 0.15 + 0.03 * len(text)
        height = 0.3
        if idx > 0 and prev_block:
            if set(block).intersection(set(prev_block)):
                x_current += overlap_shift
            else:
                x_current += separate_shift
        ellipse = patches.Ellipse((x_current, 0.5), width, height,
                                  edgecolor='green', facecolor='lightgreen', lw=2)
        ax1.add_patch(ellipse)
        ax1.text(x_current, 0.5, text, fontsize=12, ha='center', va='center',
                 color='darkgreen', fontweight='bold')
        log_msg(f"(Сравнение) Максимальный блок {idx + 1}: [{text}], x={x_current:.2f}")
        prev_block = block

    ax2.set_title("Минимальное покрытие", fontsize=12)
    ax2.axis('off')
    overlap_shift2 = 0.07
    separate_shift2 = 0.22
    x_current2 = 0.1
    prev_block2 = None
    for idx, block in enumerate(min_cover):
        text = ", ".join(block)
        width = 0.15 + 0.03 * len(text)
        height = 0.3
        if idx > 0 and prev_block2:
            if set(block).intersection(set(prev_block2)):
                x_current2 += overlap_shift2
            else:
                x_current2 += separate_shift2
        ellipse = patches.Ellipse((x_current2, 0.5), width, height,
                                  edgecolor='red', facecolor='mistyrose', lw=2)
        ax2.add_patch(ellipse)
        ax2.text(x_current2, 0.5, text, fontsize=12, ha='center', va='center',
                 color='darkred', fontweight='bold')
        log_msg(f"(Сравнение) Минимальный блок {idx + 1}: [{text}], x={x_current2:.2f}")
        prev_block2 = block

    plt.tight_layout()
    output_path = "comparison_coverings.png"
    plt.savefig(output_path)
    log_msg(f"Сравнительное покрытие сохранено в файл: {output_path}")
    plt.close(fig)


# ------------------- Экспорт бинарной матрицы в CSV ------------------- #
def export_matrix_csv():
    global LAST_DF_UPPER
    if LAST_DF_UPPER is not None:
        LAST_DF_UPPER.to_csv("binary_matrix.csv", index=True)
        log_msg("Бинарная матрица экспортирована в файл binary_matrix.csv")
    else:
        log_msg("Нет данных для экспорта бинарной матрицы.")


def draw_compatibility_graph(automata, bin_matrix):
    G = nx.DiGraph()
    compatible_pairs = []
    for (s0, s1), val in bin_matrix.items():
        if val == 1:
            compatible_pairs.append((s0, s1))

    for (p1, p2) in compatible_pairs:
        G.add_node(f"{p1}{p2}")

    for (p1, p2) in compatible_pairs:
        for inp in automata.alphabet:
            next1 = automata.table[p1][inp][0]
            next2 = automata.table[p2][inp][0]
            if next1 == '-' or next2 == '-':
                continue
            n1, n2 = min(next1, next2, key=int), max(next1, next2, key=int)
            if bin_matrix.get((n1, n2), 0) == 1:
                G.add_edge(f"{p1}{p2}", f"{n1}{n2}")

    fig, ax = plt.subplots(figsize=(8, 6))
    ax.set_title("Граф совместимых пар состояний", fontsize=14)
    pos = nx.spring_layout(G, k=0.8)
    nx.draw_networkx_nodes(G, pos, node_color='lightblue', node_size=600, ax=ax, edgecolors='black')
    nx.draw_networkx_labels(G, pos, font_color='darkblue', font_weight='bold', ax=ax)
    nx.draw_networkx_edges(G, pos, edge_color='gray', arrows=True, ax=ax, arrowstyle='-|>')
    plt.axis('off')
    plt.tight_layout()
    plt.savefig("compatibility_graph.png")
    plt.close(fig)
    print("Граф совместимых пар сохранён в файл: compatibility_graph.png")


# ===================== Алгоритмическая часть (Anger-Pohl) =====================
def anger_pohl(automata):
    log_msg(">>> Запуск алгоритма Anger-Pohl")
    aut = deepcopy(automata)
    blocks_row = defaultdict(set)
    binMatrix = {}

    log_msg("Формирование бинарной матрицы (выявление несовместимых пар)...")
    for s0, a0 in aut.table.items():
        for s1, a1 in aut.table.items():
            if s0 == s1:
                continue
            min_s = min(s0, s1)
            max_s = max(s0, s1)
            key, res = calculate(min_s, max_s, aut)
            global way
            way = set()
            binMatrix[key] = res
            log_msg(f"Пара ({min_s}, {max_s}): совместимость = {res}")
            if res:
                blocks_row[min_s].add(max_s)
    log_msg("Бинарная матрица успешно сформирована.")

    global LAST_BIN_MATRIX
    LAST_BIN_MATRIX = binMatrix

    log_msg("Поиск максимальных блоков покрытия...")
    res_list = []
    for i, a in blocks_row.items():
        for subset in is_block(sorted(a), binMatrix):
            res_list.append([i] + list(subset))
            log_msg(f"Найден блок: {[i] + list(subset)}")
    max_cover = []
    for cb in res_list:
        if not any(set(cb).issubset(set(other)) and cb != other for other in res_list):
            max_cover.append(sorted(cb, key=lambda x: int(x)))
    log_msg(f"Максимальное покрытие: {max_cover}")
    visualization(max_cover, binMatrix, "max")
    draw_coverings_with_overlap(max_cover, "max_cover")

    log_msg("Запуск минимизации покрытия методом set cover...")
    min_cover = minimize_cover(max_cover, aut)
    log_msg(f"Минимальное покрытие: {min_cover}")
    visualization(min_cover, binMatrix, "min")
    draw_coverings_with_overlap(min_cover, "min_cover")

    global LAST_AUTOMATA, LAST_MAX_COVER, LAST_MIN_COVER
    LAST_AUTOMATA = automata
    LAST_MAX_COVER = max_cover
    LAST_MIN_COVER = min_cover

    draw_comparison_coverings(max_cover, min_cover)
    generate_report_docx(automata, max_cover, min_cover)
    log_msg(">>> Алгоритм Anger-Pohl завершён.")


def minimize_cover(max_cover, automata):
    S = set(automata.states)
    candidate_blocks = [set(block) for block in max_cover]
    n = len(candidate_blocks)
    best = None
    log_msg("Начало поиска оптимального минимального покрытия...")
    for r in range(1, n + 1):
        for comb in combinations(candidate_blocks, r):
            if set().union(*comb) == S:
                best = list(comb)
                log_msg(f"Найдено покрытие из {r} блоков: {list(map(sorted, comb))}")
                break
        if best is not None:
            break
    if best is None:
        best = candidate_blocks
        log_msg("Не удалось найти оптимальное покрытие, используем все блоки.")
    final_sorted = [sorted(list(b), key=lambda x: int(x)) for b in best if b]
    final_sorted.sort(key=lambda block: int(block[0]) if block else 0)
    return final_sorted


def generate_report_docx(automata, max_cover, min_cover):
    from docx import Document
    from docx.shared import Inches

    log_msg("Начало генерации отчёта.")
    doc = Document()
    doc.add_heading("Отчёт по покрытию автомата", level=1)

    doc.add_heading("Исходная таблица автомата", level=2)
    doc.add_paragraph(str(automata.table))

    doc.add_heading("Максимальное покрытие", level=2)
    for i, block in enumerate(max_cover):
        doc.add_paragraph(f"Блок {i + 1}: {block}")

    doc.add_heading("Минимальное покрытие", level=2)
    for i, block in enumerate(min_cover):
        doc.add_paragraph(f"Блок {i + 1}: {block}")

    doc.add_heading("Визуализация бинарной матрицы", level=2)
    doc.add_picture("triangular_blocks_and_matrix_max.png", width=Inches(6))

    doc.add_heading("Логи хода решения", level=2)
    doc.add_paragraph("\n".join(LOGS))

    report_filename = "coverage_report.docx"
    doc.save(report_filename)
    log_msg(f"Отчёт сохранён в файл {report_filename}")

    pd.DataFrame({"Логи": LOGS}).to_excel("logs.xlsx", index=False)
    log_msg("Логи сохранены в файл logs.xlsx")


def visualization(blocks, bin_matrix, identifier=""):
    log_msg(f"Построение визуализации бинарной матрицы ({identifier})...")
    vertices = sorted({s for block in blocks for s in block}, key=lambda x: int(x) if x.isdigit() else x)
    idx_map = {s: i for i, s in enumerate(vertices)}
    matrix = np.full((len(vertices), len(vertices)), '', dtype=object)

    for (i, j), val in bin_matrix.items():
        if i in idx_map and j in idx_map:
            row, col = idx_map[i], idx_map[j]
            if row < col:
                matrix[row][col] = str(val)
    df_upper = pd.DataFrame(matrix, index=vertices, columns=vertices)

    global LAST_DF_UPPER
    LAST_DF_UPPER = df_upper.copy()

    fig, ax = plt.subplots(figsize=(10, 10))
    ax.axis('off')
    ax.set_title(f"Визуализация бинарной матрицы", fontsize=16, pad=20)

    table_plot = ax.table(cellText=df_upper.values,
                          rowLabels=df_upper.index,
                          colLabels=df_upper.columns,
                          loc='center',
                          cellLoc='center')
    table_plot.auto_set_font_size(False)
    table_plot.set_fontsize(10)
    table_cells = table_plot.get_celld()

    for (row, col), cell in table_cells.items():
        if row == 0 and col >= 0:
            cell.set_facecolor('#FAFAD2')
            cell.set_edgecolor('#666666')
            cell.set_text_props(color='#111111', fontweight='bold')
        elif col == -1 and row >= 0:
            cell.set_facecolor('#F0E68C')
            cell.set_edgecolor('#666666')
            cell.set_text_props(color='#111111', fontweight='bold')
        else:
            txt = cell.get_text().get_text().strip()
            if row == col and row != 0:
                cell.set_facecolor('#E6E6FA')
            elif row % 2 == 0:
                cell.set_facecolor('#e0f7fa')
            else:
                cell.set_facecolor('#ffffff')

            if txt == "1":
                cell.set_facecolor('#d0f0c0')
            elif txt == "0":
                cell.set_facecolor('#f0d0d0')

            cell.set_edgecolor('#999999')

    plt.tight_layout()
    output_path = f"triangular_blocks_and_matrix_{identifier}.png"
    plt.savefig(output_path)
    log_msg(f"Визуализация сохранена в файл: {output_path}")
    ods_path = f"triangular_blocks_and_matrix_{identifier}.xlsx"
    df_upper.to_excel(ods_path)
    log_msg(f"Данные матрицы сохранены в файл: {ods_path}")
    plt.close(fig)


def calculate(s0_, s1_, aut_):
    s0 = deepcopy(s0_)
    s1 = deepcopy(s1_)
    aut = deepcopy(aut_)
    a0 = aut.table[s0]
    a1 = aut.table[s1]
    log_msg(f"Вычисление совместимости для {s0} и {s1} ...")

    for inp in aut.alphabet:
        if a0[inp][1] != a1[inp][1] and a0[inp][1] != "-" and a1[inp][1] != "-":
            log_msg(f"Несовместимость по символу '{inp}' для {s0} и {s1}")
            return (min(s0, s1), max(s0, s1)), 0

    Yav_Soot = True
    for inp in aut.alphabet:
        for i in range(2):
            if not (a0[inp][i] == a1[inp][i] or a0[inp][i] == "-" or a1[inp][i] == "-"):
                Yav_Soot = False
                log_msg(f"Отличие (позиция {i}) по '{inp}' для {s0} и {s1}")
                break
        if not Yav_Soot:
            break

    if Yav_Soot:
        log_msg(f"Состояния {s0} и {s1} явно совместимы.")
        return (min(s0, s1), max(s0, s1)), 1

    global way
    if (min(s0, s1), max(s0, s1)) in way:
        log_msg(f"Состояния {s0} и {s1} уже проверялись – считаем совместимыми.")
        return (min(s0, s1), max(s0, s1)), 1
    else:
        way.add((min(s0, s1), max(s0, s1)))

    coord = get_way(a0, a1)
    log_msg(f"Переходы для {s0} и {s1}: {coord}")
    ans = []
    for c in coord:
        res = calculate(*c, aut)[1]
        ans.append(res)
        log_msg(f"Рекурсивное вычисление для перехода {c}: {res}")
    final_val = int(all(ans))
    log_msg(f"Итоговая совместимость для {s0} и {s1}: {final_val}")
    return (min(s0, s1), max(s0, s1)), final_val


def get_way(a0_, a1_):
    a0 = deepcopy(a0_)
    a1 = deepcopy(a1_)
    ans = []
    for inp, res in a0.items():
        if res[0] != '-' and a1[inp][0] != '-' and res[0] != a1[inp][0]:
            ans.append((min(res[0], a1[inp][0]), max(res[0], a1[inp][0])))
    return ans


def is_block(block, binMatrix):
    incor = [pair for pair in combinations(block, 2) if binMatrix.get(pair, 0) != 1]
    if not incor:
        log_msg(f"Все пары в блоке {block} совместимы.")
        return [block]
    max_blocks = set()
    for s in block:
        temp = block.copy()
        temp.remove(s)
        for sub in is_block(temp, binMatrix):
            if sub:
                max_blocks.add(tuple(sub))
    final_blocks = []
    for candidate in max_blocks:
        if not any(set(candidate).issubset(set(other)) and candidate != other for other in max_blocks):
            final_blocks.append(candidate)
    log_msg(f"Подблоки для {block}: {final_blocks}")
    return final_blocks


# -------------------- Классы автомата -------------------- #
class BaseAutomata:
    def __init__(self, states, initial_state, alphabet):
        if initial_state not in states:
            raise ValueError("Начальное состояние должно входить в множество состояний")
        self.states = states
        self.state = initial_state
        self.alphabet = alphabet


class MealyAutomata(BaseAutomata):
    def __init__(self, states, initial_state, alphabet, table):
        super().__init__(states, initial_state, alphabet)
        self.table = table
        self.state = initial_state

    def step(self, inp):
        self.state, reaction = self.table[self.state][inp]
        return reaction


# -------------------- Статистика + Формула-калькулятор -------------------- #
def show_statistics():
    global LAST_AUTOMATA, LAST_MAX_COVER, LAST_MIN_COVER
    if LAST_AUTOMATA is None or LAST_MAX_COVER is None or LAST_MIN_COVER is None:
        return "<html><body><p>Нет данных для статистики.</p></body></html>"
    num_states = len(LAST_AUTOMATA.states)
    num_max_blocks = len(LAST_MAX_COVER)
    num_min_blocks = len(LAST_MIN_COVER)
    html = f"""
    <html>
      <head>
        <style>
          body {{font-family: Arial, sans-serif; background-color: #ffffff; padding: 10px; color: #333;}}
          table {{border-collapse: collapse; width: 100%; margin-top: 10px;}}
          th, td {{border: 1px solid #999; padding: 8px; text-align: center;}}
          th {{background-color: #ddd;}}
        </style>
      </head>
      <body>
        <h2>Статистика автомата</h2>
        <table>
          <tr>
            <th>Показатель</th>
            <th>Значение</th>
          </tr>
          <tr>
            <td>Количество состояний</td>
            <td>{num_states}</td>
          </tr>
          <tr>
            <td>Блоков (макс. покрытие)</td>
            <td>{num_max_blocks}</td>
          </tr>
          <tr>
            <td>Блоков (мин. покрытие)</td>
            <td>{num_min_blocks}</td>
          </tr>
        </table>
        <h3>Максимальное покрытие</h3>
        <p>{LAST_MAX_COVER}</p>
        <h3>Минимальное покрытие</h3>
        <p>{LAST_MIN_COVER}</p>
      </body>
    </html>
    """
    return html


# -------------------- Плавающая кнопка (FAB) -------------------- #
class FloatingActionButton(QtWidgets.QToolButton):
    """
    Плавающая круглая кнопка, которая открывает меню при клике.
    """
    def __init__(self, icon_path=None, parent=None):
        super().__init__(parent)
        self.setIcon(QtGui.QIcon(icon_path) if icon_path and os.path.exists(icon_path) else QtGui.QIcon())
        self.setIconSize(QtCore.QSize(32, 32))
        self.setFixedSize(QtCore.QSize(50, 50))
        self.setStyleSheet("""
            QToolButton {
                border-radius: 25px;
                background-color: #ff4081;
                color: white;
                font-weight: bold;
            }
            QToolButton:hover {
                background-color: #f50057;
            }
        """)
        self.setToolButtonStyle(QtCore.Qt.ToolButtonTextUnderIcon)
        self.setPopupMode(QtWidgets.QToolButton.InstantPopup)


# -------------------- Основной GUI (PyQt5) -------------------- #
class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Программа: Максимальное и минимальное покрытия (Enhanced)")
        self.resize(1100, 850)

        self.menu_bar = self.menuBar()
        theme_menu = self.menu_bar.addMenu("Theme")
        self.dark_action = QtWidgets.QAction("Dark Mode", self, checkable=True)
        self.dark_action.triggered.connect(self.toggleDarkMode)
        theme_menu.addAction(self.dark_action)

        QtWidgets.QApplication.setStyle("Fusion")
        self.setStyleSheet("""
            QMainWindow {
                background: qlineargradient(x1: 0, y1: 0, x2: 1, y2: 1,
                                           stop: 0 #E8E8E8, stop: 1 #B0C4DE);
            }
            QTabWidget::pane {
                background: #ffffff;
                border: 1px solid #aaa;
                border-radius: 5px;
            }
            QTabBar::tab {
                background: #f0f0f0;
                color: #000;
                padding: 8px 16px;
                margin-right: 4px;
                border-radius: 5px 5px 0 0;
            }
            QTabBar::tab:selected {
                background: #ffffff;
                font-weight: bold;
                border-bottom: 2px solid #42a3f5;
            }
            QTableWidget {
                background-color: #ffffff;
                gridline-color: #ccc;
                alternate-background-color: #f7f7f7;
            }
            QTableWidget::item:hover {
                background-color: #e3f2fd;
            }
            QHeaderView::section {
                background-color: #fafafa;
                color: #333;
                font-weight: bold;
                border: 1px solid #ccc;
            }
            QPushButton {
                background-color: #eaeaea;
                color: #000;
                border: 1px solid #ccc;
                border-radius: 4px;
                padding: 6px 12px;
            }
            QPushButton:hover {
                background-color: #dadada;
            }
            QLineEdit {
                background-color: #ffffff;
                color: #000;
                border: 1px solid #ccc;
                padding: 4px;
            }
            QLabel {
                color: #333;
            }
            QTextEdit, QTextBrowser {
                background-color: #ffffff;
                color: #000;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QProgressBar {
                background-color: #f0f0f0;
                border: 1px solid #aaa;
                border-radius: 5px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #42a3f5;
            }
        """)

        self.tabs = QtWidgets.QTabWidget()
        self.setCentralWidget(self.tabs)

        self.input_tab = QtWidgets.QWidget()
        self.setupInputTab()
        self.tabs.addTab(self.input_tab, "Ввод")

        self.result_tab = QtWidgets.QWidget()
        self.setupResultTab()
        self.tabs.addTab(self.result_tab, "Результат")

        self.calc_tab = QtWidgets.QWidget()
        self.setupCalcTab()
        self.tabs.addTab(self.calc_tab, "Калькулятор")

        self.progress_bar = QtWidgets.QProgressBar()
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setVisible(False)
        self.statusBar().addWidget(self.progress_bar)

        self.fab = FloatingActionButton("fab_icon.png", self)
        self.fab.setToolTip("Дополнительные действия")
        fab_menu = QtWidgets.QMenu(self)
        action_logs = fab_menu.addAction(QtGui.QIcon("action1_icon.png") if os.path.exists("action1_icon.png") else QtGui.QIcon(),
                                         "Показать логи", self.showLogs)
        action_logs.setToolTip("Посмотреть подробные логи вычислений")
        action_cover = fab_menu.addAction(QtGui.QIcon("action2_icon.png") if os.path.exists("action2_icon.png") else QtGui.QIcon(),
                                          "Покрытия", self.showCoverings)
        action_cover.setToolTip("Посмотреть графическое представление покрытий")
        action_graph = fab_menu.addAction(QtGui.QIcon("action3_icon.png") if os.path.exists("action3_icon.png") else QtGui.QIcon(),
                                          "Граф совместимых пар", self.showCompatibilityGraph)
        action_graph.setToolTip("Построить граф совместимых пар состояний")

        self.fab.setMenu(fab_menu)

        # Разместим FAB в правом нижнем углу
        self.fab_container = QtWidgets.QWidget(self)
        self.fab_layout = QtWidgets.QVBoxLayout(self.fab_container)
        self.fab_layout.addWidget(self.fab, alignment=QtCore.Qt.AlignRight | QtCore.Qt.AlignBottom)
        self.fab_container.setLayout(self.fab_layout)
        self.fab_container.setAttribute(QtCore.Qt.WA_TransparentForMouseEvents, False)
        self.setCorner(QtCore.Qt.BottomRightCorner, QtCore.Qt.RightDockWidgetArea)
        self.fab_container.setGeometry(self.width()-80, self.height()-130, 50, 50)
        self.fab.raise_()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.fab_container.setGeometry(self.width()-80, self.height()-130, 50, 50)

    def setupInputTab(self):
        """
        Добавляем спинбоксы для выбора числа состояний и столбцов.
        Кнопка 'Сгенерировать таблицу' создаёт новую пустую таблицу.
        """
        layout = QtWidgets.QVBoxLayout(self.input_tab)

        # Блок для выбора размеров таблицы
        dim_layout = QtWidgets.QHBoxLayout()
        dim_layout.addWidget(QtWidgets.QLabel("Число состояний:"))
        self.num_states_spin = QtWidgets.QSpinBox()
        self.num_states_spin.setRange(1, 50)
        self.num_states_spin.setValue(9)  # по умолчанию
        dim_layout.addWidget(self.num_states_spin)

        dim_layout.addWidget(QtWidgets.QLabel("Число столбцов (входов):"))
        self.num_cols_spin = QtWidgets.QSpinBox()
        self.num_cols_spin.setRange(1, 26)
        self.num_cols_spin.setValue(2)  # по умолчанию (a, b)
        dim_layout.addWidget(self.num_cols_spin)

        self.gen_table_btn = QtWidgets.QPushButton("Сгенерировать таблицу")
        self.gen_table_btn.clicked.connect(self.generateTable)
        dim_layout.addWidget(self.gen_table_btn)

        layout.addLayout(dim_layout)

        # Информация над таблицей
        top_layout = QtWidgets.QHBoxLayout()
        instr_label = QtWidgets.QLabel(
            "<html><head><style>"
            "p { color: #444; font-size: 14px; line-height: 1.4; }"
            "</style></head><body>"
            "<p><b>Формат ячеек</b>: <i>состояние, реакция</i>, напр. <i>5, -</i> или <i>3, x</i>.</p>"
            "</body></html>"
        )
        instr_label.setWordWrap(True)
        top_layout.addWidget(instr_label)

        self.gif_label = QtWidgets.QLabel()
        if os.path.exists("loading.gif"):
            movie = QtGui.QMovie("loading.gif")
            self.gif_label.setMovie(movie)
            movie.start()
        top_layout.addWidget(self.gif_label, alignment=QtCore.Qt.AlignRight | QtCore.Qt.AlignTop)
        layout.addLayout(top_layout)

        # Создаём таблицу (по умолчанию 9 строк, 2 столбца + 1 для State)
        self.table_widget = QtWidgets.QTableWidget()
        layout.addWidget(self.table_widget)

        # Кнопка "Вычислить покрытие"
        self.compute_button = QtWidgets.QPushButton("Вычислить покрытие")
        self.compute_button.clicked.connect(self.computeCoverage)
        layout.addWidget(self.compute_button)

        # При первом запуске заполняем дефолтную (9x2) таблицу:
        self.generateTable()

    def generateTable(self):
        """
        Генерирует новую таблицу, исходя из значения спинбоксов.
        Первая колонка: State (только чтение).
        Остальные: входные символы (a,b,c,...).
        Ячейки изначально пустые, чтобы пользователь заполнял "состояние, реакция".
        """
        row_count = self.num_states_spin.value()
        col_count = self.num_cols_spin.value()

        self.table_widget.clear()
        self.table_widget.setRowCount(row_count)
        self.table_widget.setColumnCount(col_count + 1)

        # Задаём имена столбцов
        alphabet = [chr(ord('a') + i) for i in range(col_count)]
        headers = ["State"] + alphabet
        self.table_widget.setHorizontalHeaderLabels(headers)
        self.table_widget.verticalHeader().setVisible(False)
        self.table_widget.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch)

        # Заполняем первый столбец "State" (read-only)
        for i in range(row_count):
            state_item = QtWidgets.QTableWidgetItem(str(i + 1))
            state_item.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEnabled)
            self.table_widget.setItem(i, 0, state_item)

    def setupResultTab(self):
        layout = QtWidgets.QVBoxLayout(self.result_tab)

        self.view_logs_button = QtWidgets.QPushButton("Просмотр логов")
        self.view_logs_button.clicked.connect(self.showLogs)
        layout.addWidget(self.view_logs_button)

        self.view_matrix_button = QtWidgets.QPushButton("Просмотр бинарной матрицы")
        self.view_matrix_button.clicked.connect(self.showMatrixImage)
        layout.addWidget(self.view_matrix_button)

        self.view_coverings_button = QtWidgets.QPushButton("Просмотр покрытий")
        self.view_coverings_button.clicked.connect(self.showCoverings)
        layout.addWidget(self.view_coverings_button)

        self.compare_coverings_button = QtWidgets.QPushButton("Сравнение покрытий")
        self.compare_coverings_button.clicked.connect(self.showComparisonCoverings)
        layout.addWidget(self.compare_coverings_button)

        self.show_stats_button = QtWidgets.QPushButton("Показать статистику")
        self.show_stats_button.clicked.connect(self.showStatistics)
        layout.addWidget(self.show_stats_button)

        self.export_csv_button = QtWidgets.QPushButton("Экспорт бинарной матрицы в CSV")
        self.export_csv_button.clicked.connect(export_matrix_csv)
        layout.addWidget(self.export_csv_button)

        self.result_text = QtWidgets.QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(self.result_text)

        self.open_report_button = QtWidgets.QPushButton("Открыть отчёт (coverage_report.docx)")
        self.open_report_button.clicked.connect(self.openReport)
        layout.addWidget(self.open_report_button)

        self.view_compatible_graph_button = QtWidgets.QPushButton("Построить граф совместимых пар")
        self.view_compatible_graph_button.clicked.connect(self.showCompatibilityGraph)
        layout.addWidget(self.view_compatible_graph_button)

    def setupCalcTab(self):
        layout = QtWidgets.QVBoxLayout(self.calc_tab)

        formula_label = QtWidgets.QLabel(
            "<html><body>"
            "<p style='font-size:14pt; color:#333;'>"
            "Проверка неравенства:"
            "<br>||G|| - ||P<sub>max</sub>|| &le; min(||S||, ||G<sub>max</sub>||) - ||G||"
            "</p></body></html>"
        )
        layout.addWidget(formula_label)

        input_layout = QtWidgets.QGridLayout()
        input_layout.addWidget(QtWidgets.QLabel("||G||:"), 0, 0)
        self.g_edit = QtWidgets.QLineEdit()
        input_layout.addWidget(self.g_edit, 0, 1)

        input_layout.addWidget(QtWidgets.QLabel("||Pmax||:"), 1, 0)
        self.pmax_edit = QtWidgets.QLineEdit()
        input_layout.addWidget(self.pmax_edit, 1, 1)

        input_layout.addWidget(QtWidgets.QLabel("||S||:"), 2, 0)
        self.s_edit = QtWidgets.QLineEdit()
        input_layout.addWidget(self.s_edit, 2, 1)

        input_layout.addWidget(QtWidgets.QLabel("||Gmax||:"), 3, 0)
        self.gmax_edit = QtWidgets.QLineEdit()
        input_layout.addWidget(self.gmax_edit, 3, 1)

        layout.addLayout(input_layout)

        self.calc_button = QtWidgets.QPushButton("Рассчитать")
        self.calc_button.clicked.connect(self.checkInequality)
        layout.addWidget(self.calc_button)

        self.calc_result = QtWidgets.QTextEdit()
        self.calc_result.setReadOnly(True)
        layout.addWidget(self.calc_result)

    def computeCoverage(self):
        global LOGS
        LOGS = []
        input_table = {}
        states = []

        # Собираем алфавит по числу столбцов (кроме первого)
        col_count = self.table_widget.columnCount() - 1
        alphabet = [chr(ord('a') + i) for i in range(col_count)]

        def parse_cell(item, row_idx, col_name):
            if item is None or not item.text().strip():
                QtWidgets.QMessageBox.warning(
                    self,
                    "Ошибка ввода",
                    f"Пустое или некорректное значение в строке {row_idx+1}, столбец '{col_name}'.\n"
                    "Формат: 'состояние, реакция' (например: '5, x')."
                )
                raise ValueError("Некорректный формат ячейки")
            text = item.text().strip()
            parts = [p.strip() for p in text.split(",")]
            if len(parts) < 2:
                QtWidgets.QMessageBox.warning(
                    self,
                    "Ошибка ввода",
                    f"Некорректный формат в строке {row_idx+1}, столбец '{col_name}'. "
                    "Ожидается два элемента, разделённые запятой."
                )
                raise ValueError("Некорректный формат ячейки")
            return parts[:2]

        try:
            row_count = self.table_widget.rowCount()
            for i in range(row_count):
                state_item = self.table_widget.item(i, 0)
                if not state_item or not state_item.text().strip():
                    continue

                st = state_item.text().strip()
                states.append(st)
                input_table[st] = {}

                for c_i, col_sym in enumerate(alphabet, start=1):
                    cell_item = self.table_widget.item(i, c_i)
                    parsed = parse_cell(cell_item, i, col_sym)
                    input_table[st][col_sym] = parsed

            if not states:
                QtWidgets.QMessageBox.warning(self, "Ошибка ввода", "Таблица не содержит состояний.")
                return

            log_msg("Начато чтение таблицы из интерфейса.")
            self.progress_bar.setVisible(True)
            QtWidgets.QApplication.processEvents()

            automata = MealyAutomata(states, states[0], alphabet, input_table)
            log_msg("Таблица автомата успешно считана.")

            anger_pohl(automata)

            self.progress_bar.setVisible(False)
            self.result_text.setPlainText("\n".join(LOGS))
            QtWidgets.QMessageBox.information(
                self, "Выполнено",
                "Расчёты завершены. Отчёт сохранён в 'coverage_report.docx'."
            )
            self.tabs.setCurrentWidget(self.result_tab)

        except ValueError:
            self.progress_bar.setVisible(False)
            return
        except Exception as e:
            self.progress_bar.setVisible(False)
            QtWidgets.QMessageBox.critical(
                self, "Ошибка", f"Произошла непредвиденная ошибка:\n{str(e)}"
            )

    def showLogs(self):
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Подробные логи решения")
        dlg.resize(700, 500)
        layout = QtWidgets.QVBoxLayout(dlg)
        text_browser = QtWidgets.QTextBrowser()
        html = """
        <html>
          <head>
            <style>
              body {
                font-family: 'Segoe UI', sans-serif;
                font-size: 12pt;
                background-color: #fff;
                color: #333;
                padding: 10px;
              }
              .log {
                border-bottom: 1px solid #ccc;
                padding: 5px;
                margin-bottom: 5px;
              }
            </style>
          </head>
          <body>
        """
        for line in LOGS:
            html += f"<div class='log'>{line}</div>"
        html += "</body></html>"
        text_browser.setHtml(html)
        layout.addWidget(text_browser)
        dlg.exec_()

    def showMatrixImage(self):
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Бинарная матрица")
        dlg.resize(800, 800)
        layout = QtWidgets.QVBoxLayout(dlg)
        label = QtWidgets.QLabel()
        pixmap = QtGui.QPixmap("triangular_blocks_and_matrix_max.png")
        if pixmap.isNull():
            label.setText("Изображение не найдено (triangular_blocks_and_matrix_max.png).")
        else:
            label.setPixmap(pixmap.scaled(dlg.size(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation))
        layout.addWidget(label)
        dlg.exec_()

    def showCoverings(self):
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Покрытия автомата (максимальное)")
        dlg.resize(800, 400)
        layout = QtWidgets.QVBoxLayout(dlg)
        label = QtWidgets.QLabel()
        pixmap = QtGui.QPixmap("coverings_max_cover.png")
        if pixmap.isNull():
            label.setText("Изображение покрытий (coverings_max_cover.png) не найдено.")
        else:
            label.setPixmap(pixmap.scaled(dlg.size(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation))
        layout.addWidget(label)
        dlg.exec_()

    def showComparisonCoverings(self):
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Сравнительное покрытие автомата")
        dlg.resize(900, 500)
        layout = QtWidgets.QVBoxLayout(dlg)
        label = QtWidgets.QLabel()
        pixmap = QtGui.QPixmap("comparison_coverings.png")
        if pixmap.isNull():
            label.setText("Изображение сравнительного покрытия (comparison_coverings.png) не найдено.")
        else:
            label.setPixmap(pixmap.scaled(dlg.size(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation))
        layout.addWidget(label)
        dlg.exec_()

    def showStatistics(self):
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Статистика покрытия")
        dlg.resize(600, 400)
        layout = QtWidgets.QVBoxLayout(dlg)
        text_browser = QtWidgets.QTextBrowser()
        html = show_statistics()
        text_browser.setHtml(html)
        layout.addWidget(text_browser)
        dlg.exec_()

    def openReport(self):
        report_file = "coverage_report.docx"
        if os.path.exists(report_file):
            if sys.platform.startswith("darwin"):
                subprocess.call(("open", report_file))
            elif sys.platform.startswith("linux"):
                subprocess.call(("xdg-open", report_file))
            elif sys.platform.startswith("win"):
                os.startfile(report_file)
        else:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Файл отчёта не найден.")

    def showCompatibilityGraph(self):
        from PyQt5.QtWidgets import QMessageBox

        if LAST_BIN_MATRIX is None or LAST_AUTOMATA is None:
            QMessageBox.warning(
                self,
                "Данные отсутствуют",
                "Сначала вычислите покрытие, чтобы получить автомат и бинарную матрицу."
            )
            return

        draw_compatibility_graph(LAST_AUTOMATA, LAST_BIN_MATRIX)

        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Граф совместимых пар состояний")
        dlg.resize(800, 600)
        layout = QtWidgets.QVBoxLayout(dlg)
        label = QtWidgets.QLabel()
        pixmap = QtGui.QPixmap("compatibility_graph.png")
        if pixmap.isNull():
            label.setText("Изображение (compatibility_graph.png) не найдено.")
        else:
            label.setPixmap(pixmap.scaled(dlg.size(), QtCore.Qt.KeepAspectRatio, QtCore.Qt.SmoothTransformation))
        layout.addWidget(label)
        dlg.exec_()

    def checkInequality(self):
        g_str = self.g_edit.text().strip()
        pmax_str = self.pmax_edit.text().strip()
        s_str = self.s_edit.text().strip()
        gmax_str = self.gmax_edit.text().strip()

        try:
            G = int(g_str)
            Pmax = int(pmax_str)
            S = int(s_str)
            Gmax = int(gmax_str)
        except ValueError:
            self.calc_result.setPlainText("Ошибка: введите целые числа.")
            return

        lhs = G - Pmax
        rhs = min(S, Gmax) - G

        details = [
            f"||G|| = {G}",
            f"||Pmax|| = {Pmax}",
            f"||S|| = {S}",
            f"||Gmax|| = {Gmax}",
            f"left = ||G|| - ||Pmax|| = {G} - {Pmax} = {lhs}",
            f"right = min(||S||, ||Gmax||) - ||G|| = min({S}, {Gmax}) - {G} = {min(S, Gmax)} - {G} = {rhs}"
        ]

        if lhs <= rhs:
            details.append("Результат: неравенство выполняется.")
        else:
            details.append("Результат: неравенство НЕ выполняется.")

        self.calc_result.setPlainText("\n".join(details))

    def toggleDarkMode(self, checked):
        if checked:
            self.setStyleSheet("""
                QMainWindow {
                    background: #2b2b2b;
                }
                QTabWidget::pane {
                    background: #3b3b3b;
                    border: 1px solid #777;
                    border-radius: 5px;
                }
                QTabBar::tab {
                    background: #4c4c4c;
                    color: #fff;
                    padding: 8px 16px;
                    margin-right: 4px;
                    border-radius: 5px 5px 0 0;
                }
                QTabBar::tab:selected {
                    background: #5c5c5c;
                    font-weight: bold;
                    border-bottom: 2px solid #88c0d0;
                }
                QTableWidget {
                    background-color: #3b3b3b;
                    color: #eee;
                    gridline-color: #666;
                    alternate-background-color: #4c4c4c;
                }
                QTableWidget::item:hover {
                    background-color: #616161;
                }
                QHeaderView::section {
                    background-color: #444;
                    color: #eee;
                    font-weight: bold;
                    border: 1px solid #666;
                }
                QPushButton {
                    background-color: #555;
                    color: #eee;
                    border: 1px solid #666;
                    border-radius: 4px;
                    padding: 6px 12px;
                }
                QPushButton:hover {
                    background-color: #666;
                }
                QLineEdit {
                    background-color: #3b3b3b;
                    color: #fff;
                    border: 1px solid #666;
                    padding: 4px;
                }
                QLabel {
                    color: #ddd;
                }
                QTextEdit, QTextBrowser {
                    background-color: #3b3b3b;
                    color: #fff;
                    border: 1px solid #666;
                    border-radius: 4px;
                }
                QProgressBar {
                    background-color: #444;
                    border: 1px solid #666;
                    border-radius: 5px;
                    text-align: center;
                    color: #fff;
                }
                QProgressBar::chunk {
                    background-color: #88c0d0;
                }
                QToolButton {
                    border-radius: 25px;
                    background-color: #ff4081;
                    color: white;
                    font-weight: bold;
                }
                QToolButton:hover {
                    background-color: #f50057;
                }
            """)
        else:
            self.setStyleSheet("""
                QMainWindow {
                    background: qlineargradient(x1: 0, y1: 0, x2: 1, y2: 1,
                                               stop: 0 #E8E8E8, stop: 1 #B0C4DE);
                }
                QTabWidget::pane {
                    background: #ffffff;
                    border: 1px solid #aaa;
                    border-radius: 5px;
                }
                QTabBar::tab {
                    background: #f0f0f0;
                    color: #000;
                    padding: 8px 16px;
                    margin-right: 4px;
                    border-radius: 5px 5px 0 0;
                }
                QTabBar::tab:selected {
                    background: #ffffff;
                    font-weight: bold;
                    border-bottom: 2px solid #42a3f5;
                }
                QTableWidget {
                    background-color: #ffffff;
                    gridline-color: #ccc;
                    alternate-background-color: #f7f7f7;
                }
                QTableWidget::item:hover {
                    background-color: #e3f2fd;
                }
                QHeaderView::section {
                    background-color: #fafafa;
                    color: #333;
                    font-weight: bold;
                    border: 1px solid #ccc;
                }
                QPushButton {
                    background-color: #eaeaea;
                    color: #000;
                    border: 1px solid #ccc;
                    border-radius: 4px;
                    padding: 6px 12px;
                }
                QPushButton:hover {
                    background-color: #dadada;
                }
                QLineEdit {
                    background-color: #ffffff;
                    color: #000;
                    border: 1px solid #ccc;
                    padding: 4px;
                }
                QLabel {
                    color: #333;
                }
                QTextEdit, QTextBrowser {
                    background-color: #ffffff;
                    color: #000;
                    border: 1px solid #ccc;
                    border-radius: 4px;
                }
                QProgressBar {
                    background-color: #f0f0f0;
                    border: 1px solid #aaa;
                    border-radius: 5px;
                    text-align: center;
                }
                QProgressBar::chunk {
                    background-color: #42a3f5;
                }
                QToolButton {
                    border-radius: 25px;
                    background-color: #ff4081;
                    color: white;
                    font-weight: bold;
                }
                QToolButton:hover {
                    background-color: #f50057;
                }
            """)


def main():
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
